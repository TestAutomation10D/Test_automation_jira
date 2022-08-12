from docx import Document
from docx.shared import Inches


class DocumentCreation:
    def __init__(self):
        self.document_obj = Document()
        self.para_obj = None
        self.run_obj = None

    def create_para_object(self):
        self.para_obj = self.document_obj.add_paragraph()

    def add_run_to_para_obj(self):
        self.run_obj = self.para_obj.add_run()

    def add_text_to_run_obj(self, input_text):
        self.run_obj.add_text(input_text)

    def add_image_to_run_obj(self, image_path):
        self.run_obj.add_picture(image_path, width=Inches(6.0), height=Inches(3.3))

    def save_document(self, doc_path):
        self.document_obj.save(doc_path)

